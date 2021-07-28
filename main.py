#! Collateral_Assignments.py - a script that will read the names of personnel
#  and a list of collateral duties from separate Excel spreadsheets, and will
#  ask the user for primary and secondary assignments for each duty. Once selected
#  the system will print out an appropriate appointment letter for each duty.

import calendar
import pyinputplus as pyip
import openpyxl
import docx
import glob
import re
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import date


MAX_LINE_LENGTH = 28

def choose_workbook(message):
    """Asking the user to clarify which excel correlates to personnel or duties list"""
    files = glob.glob('*xlsx')
    print(message)
    output = pyip.inputMenu(files, numbered=True)
    return output

def get_personnel():
    """Reading the current personnel from the provided Excel sheet."""
    file = choose_workbook("Which Excel has the personnel of the department?")
    wb = openpyxl.load_workbook(file)
    sheet = wb.active
    for row in range(1, sheet.max_row + 1):
        personnel.append(sheet[f'A{row}'].value)
    wb.close()

def get_duties():
    """Reading the current duties from the provided Excel sheet."""
    file = choose_workbook("Which Excel lists the collateral duties?")
    wb = openpyxl.load_workbook(file)
    sheet = wb.active
    for row in range(1, sheet.max_row + 1):
        duties.append(sheet[f'A{row}'].value)
    wb.close()

def get_current_date():
    """Get current date with goal format of: dd MMM YYYY"""
    date_year = date.today().year
    date_month = date.today().month
    month_abbr = calendar.month_abbr[date_month]
    date_day = date.today().day
    return f'{date_day} {month_abbr} {date_year}'

def get_officers_name():
    """Getting officer's name with double-check before continuing on."""
    correct_officer = False

    while correct_officer is False:
        output = input("Enter the name of the officer who's signing the letters: ")
        print(f"You entered: {output.upper()}. Is this correct?")
        double_check = pyip.inputMenu(['Yes', 'No'], numbered=True)
        if double_check == 'Yes':
            correct_officer = True
    return output

def chose_member(message, list_of_members):
    """Present the user with a list of their personnel and return the chosen member."""
    print(message)
    output = pyip.inputMenu(list_of_members, numbered=True)
    return output

def remove_slashes(duty):
    """Removing both forward and back slashed from duty names in order to save without throwing an error."""
    output = re.sub(r'[\\/]', '-', duty)
    return output

def get_chosen_duties(duties):
    output = []
    member_is_picking_duties = True
    while member_is_picking_duties:
        print("Which duty do you need to assign?")
        duty = pyip.inputMenu(duties, numbered=True)
        print("Is this for the primary or secondary position?")
        primary_or_secondary = pyip.inputMenu(['Primary', 'Secondary'], numbered=True)
        output.append([duty, primary_or_secondary])
        print("Do you need to assign another duty?")
        more_to_do = pyip.inputMenu(["Yes", "No"], numbered=True)
        if more_to_do == "Yes":
            continue
        else:
            member_is_picking_duties = False
    return output

def split_line(duty):
    """In order to ensure formatting with long duty names, we're going to split long duty names into two lines."""
    strings = duty.split()
    first_line = strings[0]
    string_count = 1
    for i in range(1, len(strings) + 1):
        if len(first_line) + len(strings[i]) + 1 < MAX_LINE_LENGTH:
            first_line = first_line + " " + strings[i]
            string_count += 1
        else:
            break

    second_line = strings[string_count]
    if len(strings) > string_count:

        for i in range(string_count + 1, len(strings)):
            second_line = second_line + " " + strings[i]

    return f"{first_line}\n          {second_line}"

def write_letter(duty, member, holder, officer):
    """Writing the actual appointment letter with the individualized duty, member, and current date."""
    letter = docx.Document()
    styles = letter.styles
    style = styles.add_style('Times New Roman', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    date_paragraph = letter.add_paragraph(get_current_date())
    date_paragraph_format = date_paragraph.paragraph_format
    date_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_paragraph.style = letter.styles['Times New Roman']

    memorandum_paragraph = letter.add_paragraph()
    memorandum_paragraph.add_run('MEMORANDUM').bold = True
    memorandum_paragraph.style = letter.styles['Times New Roman']

    from_paragraph = letter.add_paragraph(f"From: Division Officer, Pharmacy Department, Navy Medicine Readiness and "
                                          f"Training\n           Command Lemoore\nTo:     {member}, USN\nVia:    "
                                          f"Leading Chief Petty Officer, Pharmacy Department")
    from_paragraph.style = letter.styles['Times New Roman']

    if len(duty) < MAX_LINE_LENGTH:
        sub_paragraph = letter.add_paragraph(f"Subj:  APPOINTMENT AS {holder.upper()} {duty.upper()}")
        sub_paragraph.style = letter.styles['Times New Roman']
    else:
        sub_paragraph = letter.add_paragraph(f"Subj:  APPOINTMENT AS {holder.upper()} {split_line(duty).upper()}")
        sub_paragraph.style = letter.styles['Times New Roman']

    ref_paragraph = letter.add_paragraph("Ref:    (a) NMRTC PHARMACY DEPARTMENT COLLATERAL DUTY \n\t    EXPECTATIONS\n"
                                         "           (b) NMRTC PHARMACY DEPARTMENT PERSONAL QUALIFICATION\n           "
                                         "     STANDARDS")
    ref_paragraph.style = letter.styles['Times New Roman']

    bullet_one_paragraph = letter.add_paragraph()
    bullet_one_paragraph.add_run(f"1.  Effective immediately, you are hereby appointed as the {duty} for the Pharmacy "
                                 f"Department at Navy Medicine Readiness and Training Command Lemoore.  You will be "
                                 f"guided in the conduct of your duties by reference (a) and (b) and directly "
                                 f"responsible to the Division Officer.")
    bullet_one_paragraph.style = letter.styles['Times New Roman']

    bullet_two_paragraph = letter.add_paragraph()
    bullet_two_paragraph.add_run("2.  Failure to properly execute assigned duties will result in remedial training "
                                 "from the department Leading Petty Officer for correction.  If remedial corrective "
                                 "actions prove unsuccessful, it will be the recommendation of leadership to forward "
                                 "to Leading Petty Officer, and or Senior Enlisted Leader for further disciplinary "
                                 "review and actions.")
    bullet_two_paragraph.style = letter.styles['Times New Roman']

    bullet_three_paragraph = letter.add_paragraph()
    bullet_three_paragraph.add_run("3.  This appointment will remain in effect until reassignment, unless otherwise "
                                   "directed.")
    bullet_three_paragraph.style = letter.styles['Times New Roman']

    spacing_paragraph = letter.add_paragraph()
    second_spacing_paragraph = letter.add_paragraph()
    signing_officer_paragraph = letter.add_paragraph()
    signing_officer_paragraph.add_run(f"\t\t\t\t\t\t\t{officer.upper()}")
    signing_officer_paragraph.style = letter.styles['Times New Roman']

    another_spacing_paragraph = letter.add_paragraph()
    copy_to_paragraph = letter.add_paragraph()
    copy_to_paragraph.add_run("Copy to:\nDIVO FOLDER")
    copy_to_paragraph.style = letter.styles['Times New Roman']

    letter.save(f"{remove_slashes(duty)} Appointment Letter-{holder} {member}.docx")

# Setting up blank lists to be filled in from current Excel sheets.
personnel = []
duties = []

# A list in-case there are duties that aren't yet assigned.
come_back_to_duties = []

# Populate lists from Excel.
get_personnel()
get_duties()

# Set officer's name who's signing the letters
signing_officer = get_officers_name()

# Ask if the user wants to redo the letter for every duty, or if they would like to generate just a select
# few letters.
print("Are you redoing all of the assignment letters, or just a few?")
users_choice = pyip.inputMenu(["Just a few", "All of them"], numbered=True)

if users_choice == "Just a few":
    chosen_duties = get_chosen_duties(duties)
    for duty in chosen_duties:
        write_letter(duty[0], chose_member(f"Which member do you want to assign to: {duty[0]}", personnel), duty[1],
                     signing_officer)
else:
    # Run loop until every duty, both primary and secondary, is accounted for.
    for duty in duties:
        print(f"Do you know who you want to assign as the primary holder of: {duty}?")
        check_for_assignment = pyip.inputMenu(["Yes", "No"], numbered=True)
        if check_for_assignment == "Yes":
            primary_holder = chose_member(f"Who's going to be the primary for: {duty}?", personnel)
            holder = "PRIMARY"
            write_letter(duty, primary_holder, holder, signing_officer)
        else:
            come_back_to_duties.append([duty, "PRIMARY"])

        print(f"Do you know who you want to assign as the secondary holder of: {duty}?")
        check_for_assignment = pyip.inputMenu(["Yes", "No"], numbered=True)
        if check_for_assignment == "Yes":
            primary_holder = chose_member(f"Who's going to be the secondary for: {duty}?", personnel)
            holder = "SECONDARY"
            write_letter(duty, primary_holder, holder, signing_officer)
        else:
            come_back_to_duties.append([duty, "SECONDARY"])

    # Make an Excel of duties that still need members assigned.
    duties_left_wb = openpyxl.Workbook()
    ws = duties_left_wb.active
    row_to_write_to = 2
    ws['A1'] = "Unassigned Duties"
    ws['B1'] = "Primary/Secondary"
    for entry in come_back_to_duties:
        ws[f'A{row_to_write_to}'] = entry[0]
        ws[f'B{row_to_write_to}'] = entry[1]
        row_to_write_to += 1

    duties_left_wb.save("duties_left.xlsx")
    duties_left_wb.close()
